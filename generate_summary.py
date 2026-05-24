"""
CPE232 Final Exam Summary Generator
Generates CPE232_FinalExam_Summary.docx from all 4 lecture chapters
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ─── Colors ────────────────────────────────────────────────────────────────────
DARK_BLUE   = RGBColor(0x00, 0x37, 0x6B)   # heading 1
TEAL        = RGBColor(0x00, 0x7B, 0x8A)   # heading 2
GRAY_HEADER = RGBColor(0x4A, 0x4A, 0x4A)   # table header text
GRAY_BG     = RGBColor(0xD9, 0xD9, 0xD9)   # table header bg
LIGHT_BG    = RGBColor(0xF2, 0xF2, 0xF2)   # alternate row
FONT_NAME   = "TH Sarabun New"

doc = Document()

# ─── Page Setup ────────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)

# ─── Helpers ───────────────────────────────────────────────────────────────────
def set_font(run, size=14, bold=False, italic=False, color=None):
    run.font.name      = FONT_NAME
    run.font.size      = Pt(size)
    run.font.bold      = bold
    run.font.italic    = italic
    if color:
        run.font.color.rgb = color

def add_heading1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    set_font(run, 22, bold=True, color=DARK_BLUE)
    return p

def add_heading2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    set_font(run, 18, bold=True, color=TEAL)
    return p

def add_heading3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    set_font(run, 16, bold=True, color=DARK_BLUE)
    return p

def add_body(text, size=14):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_font(run, size)
    return p

def add_bullet(text, level=0, size=13):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_font(run, size)
    return p

def add_formula(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.5)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_font(run, 13, italic=True)
    run.font.color.rgb = RGBColor(0x00, 0x37, 0x6B)
    return p

def set_cell_bg(cell, rgb: RGBColor):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    hex_color = f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def style_header_row(table):
    for cell in table.rows[0].cells:
        set_cell_bg(cell, GRAY_BG)
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.name = FONT_NAME
                run.font.size = Pt(13)

def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    # header
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_bg(hdr[i], DARK_BLUE)
        for para in hdr[i].paragraphs:
            for run in para.runs:
                run.font.bold  = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.name  = FONT_NAME
                run.font.size  = Pt(13)
    # data rows
    for r_idx, row in enumerate(rows):
        cells = t.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = val
            if r_idx % 2 == 1:
                set_cell_bg(cells[c_idx], LIGHT_BG)
            for para in cells[c_idx].paragraphs:
                for run in para.runs:
                    run.font.name = FONT_NAME
                    run.font.size = Pt(12)
    # widths
    if col_widths:
        for row in t.rows:
            for i, cell in enumerate(row.cells):
                cell.width = Cm(col_widths[i])
    return t

def page_break():
    doc.add_page_break()

def add_separator():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run("─" * 80)
    run.font.color.rgb = GRAY_BG
    run.font.size = Pt(8)

# ─── Header & Footer ───────────────────────────────────────────────────────────
def add_header_footer():
    section = doc.sections[0]
    # Header
    header = section.header
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.clear()
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = hp.add_run("CPE232 Data Models — สรุปเนื้อหาสอบปลายภาค (Chapter 6–9)")
    set_font(run, 11, italic=True, color=TEAL)
    # Footer with page number
    footer = section.footer
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.clear()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_f = fp.add_run()
    set_font(run_f, 11)
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run_f._r.append(fldChar1)
    run_f._r.append(instrText)
    run_f._r.append(fldChar2)

add_header_footer()

# ════════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ════════════════════════════════════════════════════════════════════════════════
for _ in range(5):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("สรุปเนื้อหาสอบปลายภาค")
set_font(run, 32, bold=True, color=DARK_BLUE)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("CPE232: Data Models")
set_font(run, 26, bold=True, color=TEAL)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Chapter 6: ML Classification\n"
                "Chapter 7: ML Regression\n"
                "Chapter 8: Model Parameter Tuning\n"
                "Chapter 9: ML Clustering")
set_font(run, 16, color=DARK_BLUE)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Dr. Sansiri Tarnpradab\nDepartment of Computer Engineering, KMUTT")
set_font(run, 14, italic=True)

page_break()

# ════════════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (manual)
# ════════════════════════════════════════════════════════════════════════════════
add_heading1("สารบัญ (Table of Contents)")
toc_items = [
    ("บทที่ 6: ML Classification",    "3"),
    ("บทที่ 7: ML Regression",         "7"),
    ("บทที่ 8: Model Parameter Tuning","11"),
    ("บทที่ 9: ML Clustering",         "15"),
    ("ตาราง Cheat Sheet สูตรทั้งหมด", "19"),
    ("แนวข้อสอบ - บทที่ 6",            "21"),
    ("แนวข้อสอบ - บทที่ 7",            "24"),
    ("แนวข้อสอบ - บทที่ 8",            "27"),
    ("แนวข้อสอบ - บทที่ 9",            "30"),
]
for title, pg in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"  {title}")
    set_font(run, 14)
    run2 = p.add_run(f"  .......  หน้า {pg}")
    set_font(run2, 13, italic=True)

page_break()

# ════════════════════════════════════════════════════════════════════════════════
# CHAPTER 6: ML CLASSIFICATION
# ════════════════════════════════════════════════════════════════════════════════
add_heading1("บทที่ 6: ML Classification (การจำแนกประเภทด้วย Machine Learning)")

add_heading2("6.1 ภาพรวม (Overview)")
add_body("Classification คือการทำนายว่าข้อมูลที่ป้อนเข้ามาอยู่ในกลุ่ม/ประเภทใด "
         "โดยโมเดลเรียนรู้จาก labeled data (Supervised Learning)")
add_bullet("Input: Feature vector (x₁, x₂, ..., xₙ)")
add_bullet("Output: Class label (discrete category)")
add_bullet("เปรียบเทียบกับ Regression ที่ output เป็น continuous value")

add_heading2("6.2 Decision Tree")
add_body("ต้นไม้การตัดสินใจที่แบ่งข้อมูลตาม feature ที่ให้ข้อมูลมากที่สุด")

add_heading3("Entropy (ความไม่แน่นอนของข้อมูล)")
add_formula("Entropy(t) = -Σⱼ p(j|t) · log₂ p(j|t)")
add_bullet("p(j|t) = สัดส่วนของ class j ใน node t")
add_bullet("Entropy = 0 → node บริสุทธิ์ (pure node), Entropy สูงสุด = 1 (ข้อมูลปนกันเท่ากัน)")

add_heading3("Information Gain (Gain_split)")
add_formula("Gain_split = Entropy(parent) - Σᵢ (nᵢ/n) · Entropy(child_i)")
add_bullet("nᵢ = จำนวน sample ใน child node i,   n = จำนวน sample ทั้งหมด")
add_bullet("เลือก feature ที่ให้ Gain สูงสุดมาแบ่ง node")
add_heading3("ตัวอย่าง Information Gain จากบทเรียน")
add_table(
    ["Feature", "Information Gain"],
    [
        ["age",           "0.246"],
        ["income",        "0.028"],
        ["student",       "0.131"],
        ["credit_rating", "0.047"],
    ],
    [6, 6]
)
add_body("→ เลือก age ก่อนเพราะ Gain สูงสุด (0.246)")

add_heading3("Gini Index")
add_formula("Gini(t) = 1 - Σⱼ [p(j|t)]²")
add_bullet("ใช้แทน Entropy ได้  •  Gini = 0 → pure node")

add_heading2("6.3 K-Nearest Neighbors (KNN)")
add_bullet("จำแนก sample ใหม่โดยดูจาก K เพื่อนบ้านที่ใกล้ที่สุด")
add_bullet("K น้อยเกินไป → Overfitting,  K มากเกินไป → Underfitting")
add_bullet("ต้องเลือก distance metric: Euclidean, Manhattan ฯลฯ")

add_heading2("6.4 Naive Bayes")
add_formula("P(C|X) ∝ P(C) · Π P(xᵢ|C)")
add_bullet("สมมติว่า features เป็นอิสระต่อกัน (naive assumption)")
add_bullet("ทำงานได้ดีกับข้อมูล text classification")

add_heading2("6.5 Logistic Regression")
add_formula("P(y=1|x) = 1 / (1 + e^(-(wx + b)))")
add_bullet("ใช้ Sigmoid function แปลงเป็นความน่าจะเป็น (0–1)")
add_bullet("Decision boundary: P ≥ 0.5 → class 1")

add_heading2("6.6 Support Vector Machine (SVM)")
add_bullet("หา hyperplane ที่แบ่ง class โดยมี margin สูงสุด")
add_bullet("Support Vectors = จุดข้อมูลที่ใกล้ hyperplane มากที่สุด")
add_bullet("Kernel trick: ขยายมิติข้อมูลเพื่อแก้ปัญหา non-linear")

add_heading2("6.7 Random Forest")
add_bullet("Ensemble method: รวม Decision Trees หลายต้นเข้าด้วยกัน")
add_bullet("ใช้ Bagging (Bootstrap Aggregating) → สุ่มข้อมูล + สุ่ม features")
add_bullet("ลด Overfitting ได้ดีกว่า Decision Tree เดี่ยว")

add_heading2("6.8 การประเมินโมเดล Classification")
add_heading3("Confusion Matrix")
add_table(
    ["", "Predicted: Yes", "Predicted: No"],
    [
        ["Actual: Yes", "TP (True Positive)",  "FN (Type II Error)"],
        ["Actual: No",  "FP (Type I Error)",   "TN (True Negative)"],
    ],
    [4, 5, 5]
)
add_heading3("เมตริกการประเมิน")
add_formula("Accuracy = (TP + TN) / (TP + TN + FP + FN)")
add_formula("Recall (Sensitivity) = TP / (TP + FN)")
add_formula("Precision = TP / (TP + FP)")
add_formula("F1-Score = 2 × (Precision × Recall) / (Precision + Recall)")
add_bullet("Accuracy: ความถูกต้องโดยรวม")
add_bullet("Recall: ครอบคลุม positive ได้มากแค่ไหน (สำคัญเมื่อ FN มีต้นทุนสูง เช่น โรคมะเร็ง)")
add_bullet("Precision: สิ่งที่ทำนายว่า positive นั้นถูกต้องแค่ไหน")
add_bullet("F1: ค่าเฉลี่ยแบบ harmonic mean ของ Precision และ Recall")
add_bullet("ROC Curve / AUC: วัดประสิทธิภาพโมเดล binary classification")

page_break()

# ════════════════════════════════════════════════════════════════════════════════
# CHAPTER 7: ML REGRESSION
# ════════════════════════════════════════════════════════════════════════════════
add_heading1("บทที่ 7: ML Regression (การพยากรณ์ค่าต่อเนื่อง)")

add_heading2("7.1 ภาพรวม")
add_body("Regression ทำนาย numerical value (continuous output) จาก features")
add_bullet("สมการหลัก: Data = Model + Error")
add_bullet("Independent variables (x) = Predictors")
add_bullet("Dependent variable (y) = Outcome")

add_heading2("7.2 Linear Regression")
add_formula("ŷ = w₀ + w₁x₁ + w₂x₂ + ... + wₙxₙ")
add_bullet("w₀ = intercept (bias),  wᵢ = coefficients (weights)")
add_bullet("หา w ที่ minimize loss function (MSE)")

add_heading2("7.3 Polynomial Regression")
add_formula("ŷ = w₀ + w₁x + w₂x² + ... + wₙxⁿ  (degree n > 1)")
add_bullet("ใช้เมื่อความสัมพันธ์ระหว่าง x และ y ไม่ใช่เส้นตรง")
add_bullet("degree สูงมาก → Overfitting")

add_heading2("7.4 Loss Functions")
add_heading3("Mean Squared Error (MSE)")
add_formula("MSE = (1/n) Σᵢ (yᵢ - ŷᵢ)²")
add_heading3("Root Mean Squared Error (RMSE)")
add_formula("RMSE = √[(1/n) Σᵢ (yᵢ - ŷᵢ)²]")
add_heading3("Mean Absolute Error (MAE)")
add_formula("MAE = (1/n) Σᵢ |yᵢ - ŷᵢ|")
add_table(
    ["Metric", "สูตร", "ลักษณะ"],
    [
        ["MSE",  "(1/n)Σ(y-ŷ)²",  "ให้น้ำหนัก error ใหญ่มาก (squared)"],
        ["RMSE", "√MSE",           "หน่วยเดียวกับ y, ตีความง่าย"],
        ["MAE",  "(1/n)Σ|y-ŷ|",   "ทนทานต่อ outlier กว่า MSE"],
    ],
    [3, 5, 6]
)

add_heading2("7.5 Gradient Descent")
add_body("วิธีอัปเดต weight เพื่อลด loss function ทีละ step")
add_formula("Cost function: J(w, b) = (1/n) Σᵢ (yᵢ - ŷᵢ)²")
add_formula("Update rule: w = w - α · (∂J/∂w)")
add_formula("Update rule: b = b - α · (∂J/∂b)")
add_bullet("α = Learning rate (ขนาดก้าวการเดิน)")
add_bullet("α เล็กเกินไป → เรียนช้า,  α ใหญ่เกินไป → diverge")

add_heading2("7.6 R² (Coefficient of Determination)")
add_formula("R² = 1 - (SSE / SSyy)")
add_formula("SSyy = Σ(yᵢ - ȳ)²   (Total variance)")
add_formula("SSE  = Σ(yᵢ - Ŷᵢ)²  (Unexplained variance)")
add_bullet("R² ใกล้ 1 → โมเดลอธิบายความแปรปรวนได้ดี")
add_bullet("R² = 0 → โมเดลไม่ดีกว่าการใช้ค่าเฉลี่ย")

add_heading2("7.7 Regularization (การป้องกัน Overfitting)")
add_heading3("Ridge Regression (L2)")
add_formula("Loss_Ridge = Σ(yᵢ - ŷᵢ)² + λ · Σβⱼ²")
add_bullet("เพิ่ม penalty ของ weight²  →  ลด magnitude ของ weight แต่ไม่เป็น 0")

add_heading3("Lasso Regression (L1)")
add_formula("Loss_Lasso = Σ(yᵢ - ŷᵢ)² + λ · Σ|βⱼ|")
add_bullet("เพิ่ม penalty ของ |weight|  →  บาง weight กลายเป็น 0 (feature selection)")
add_bullet("λ (lambda) = hyperparameter ควบคุมความแรงของ regularization")

add_heading2("7.8 Decision Tree Regression")
add_bullet("ใช้ Decision Tree แบ่งพื้นที่ feature space")
add_bullet("ค่าพยากรณ์ = ค่าเฉลี่ยของ sample ใน leaf node")
add_bullet("ข้อดี: ไม่ต้องสมมติความสัมพันธ์เชิงเส้น,  ข้อเสีย: Overfitting ง่าย")

page_break()

# ════════════════════════════════════════════════════════════════════════════════
# CHAPTER 8: PARAMETER TUNING
# ════════════════════════════════════════════════════════════════════════════════
add_heading1("บทที่ 8: Model Parameter Tuning (การปรับแต่งพารามิเตอร์โมเดล)")

add_heading2("8.1 Learnable Parameters vs Hyperparameters")
add_table(
    ["ประเภท", "คำอธิบาย", "ตัวอย่าง"],
    [
        ["Learnable Parameters",
         "ถูกเรียนรู้โดยโมเดลระหว่าง training โดยอัตโนมัติ (minimize loss)",
         "weights, coefficients, biases"],
        ["Hyperparameters",
         "กำหนดโดยผู้ใช้ก่อน training (ไม่ได้เรียนรู้จากข้อมูล)",
         "learning rate, max_depth, k ใน KNN"],
    ],
    [4, 7, 5]
)

add_heading2("8.2 Hyperparameters ของ Decision Tree")
add_table(
    ["Hyperparameter", "ความหมาย"],
    [
        ["max_depth",          "ความลึกสูงสุดของต้นไม้ — ป้องกัน Overfitting"],
        ["min_samples_split",  "จำนวน sample ขั้นต่ำในโหนดก่อนแยก"],
        ["min_samples_leaf",   "จำนวน sample ขั้นต่ำที่ leaf node ต้องมี"],
        ["max_features",       "จำนวน feature สูงสุดที่พิจารณาในการแยก"],
        ["criterion",          "วิธีวัดคุณภาพการแบ่ง: Entropy หรือ Gini"],
        ["splitter",           "กลยุทธ์การแบ่ง: Best split หรือ Random split"],
    ],
    [5, 9]
)

add_heading2("8.3 Overfitting และ Underfitting")
add_heading3("Overfitting")
add_bullet("โมเดลเรียนรู้ noise ใน training data มากเกินไป")
add_bullet("ลักษณะ: Training accuracy สูง, Test accuracy ต่ำ")
add_bullet("สาเหตุ: โมเดลซับซ้อนเกินไป (deep trees, too many parameters)")
add_heading3("วิธีป้องกัน Overfitting")
add_bullet("Simplify model (pruning trees, reduce parameters)")
add_bullet("Regularization (L1/L2)")
add_bullet("เพิ่มข้อมูล (More training data)")
add_bullet("Cross-Validation")
add_bullet("Dropout (Neural Networks)")
add_heading3("Underfitting")
add_bullet("โมเดลเรียบง่ายเกินไป ไม่สามารถจับ pattern ของข้อมูลได้")
add_bullet("ลักษณะ: Training และ Test accuracy ต่ำทั้งคู่")
add_heading3("วิธีป้องกัน Underfitting")
add_bullet("เพิ่มความซับซ้อนของโมเดล (Increase model complexity)")
add_bullet("Feature Engineering (เพิ่ม relevant features)")
add_bullet("ลด Regularization (Reduce regularization)")
add_bullet("เพิ่มข้อมูล (More data)")

add_heading2("8.4 Cross-Validation")
add_body("เทคนิคประเมินโมเดลโดยแบ่งข้อมูลเป็นหลายชุด training/validation")
add_table(
    ["ประเภท", "คำอธิบาย", "เหมาะกับ"],
    [
        ["k-Fold CV",
         "แบ่งข้อมูลเป็น k ส่วน วนทดสอบ k รอบ",
         "ข้อมูลทั่วไป"],
        ["Stratified k-Fold",
         "แต่ละ fold มี class distribution เหมือนชุดข้อมูลต้นฉบับ",
         "ข้อมูล imbalanced"],
        ["Leave-One-Out (LOO-CV)",
         "ใช้ 1 sample เป็น test ทีละครั้ง",
         "ข้อมูลน้อย"],
        ["Time Series CV",
         "train บน past data, test บน future data",
         "ข้อมูล time-series"],
    ],
    [4, 7, 3]
)
add_body("Validation Set vs Test Set:")
add_bullet("Validation Set: ใช้ระหว่าง training เพื่อ tune hyperparameters")
add_bullet("Test Set: ใช้ครั้งเดียวหลัง training เสร็จ เพื่อประเมิน final performance")

add_heading2("8.5 วิธีการ Hyperparameter Tuning")
add_table(
    ["วิธี", "คำอธิบาย", "ข้อดี/ข้อเสีย"],
    [
        ["Manual",        "ลองด้วยประสบการณ์และความรู้", "เข้าใจง่าย / ช้า ไม่ครอบคลุม"],
        ["Grid Search",   "ค้นหา exhaustively ทุก combination", "รับประกันว่าหา best ได้ / ใช้เวลามาก"],
        ["Random Search", "สุ่ม sample จาก parameter space", "เร็วกว่า Grid Search / อาจพลาด best"],
        ["Bayesian Opt.", "ใช้ผลก่อนหน้าเป็น prior เพื่อเลือก next point", "ฉลาด ประหยัด / ซับซ้อน"],
    ],
    [3, 6, 5]
)
add_body("Tools: scikit-optimize, Hyperopt, Optuna")

add_heading2("8.6 Gradient Descent Variants")
add_table(
    ["ประเภท", "Data ที่ใช้ต่อ update", "Update ต่อ Epoch", "ลักษณะ"],
    [
        ["Batch GD",      "ทั้งชุดข้อมูล",         "1",             "เสถียร แต่ช้า"],
        ["Stochastic GD", "1 sample",              "n (= n samples)", "เร็ว แต่ noisy"],
        ["Mini-batch GD", "batch ขนาดเล็ก (32/64)", "n/batch_size",  "สมดุลระหว่าง 2 แบบ"],
    ],
    [4, 4, 4, 4]
)

add_heading2("8.7 Challenges: Local Minima")
add_bullet("Local Minima: gradient descent หยุดที่จุดต่ำสุดใน local แต่ไม่ใช่ global minimum")
add_bullet("วิธีแก้: SGD (เพิ่ม noise), Learning Rate Scheduling, Adaptive Optimizers (Adam, RMSProp)")

page_break()

# ════════════════════════════════════════════════════════════════════════════════
# CHAPTER 9: CLUSTERING
# ════════════════════════════════════════════════════════════════════════════════
add_heading1("บทที่ 9: ML Clustering (การจัดกลุ่มข้อมูล)")

add_heading2("9.1 Unsupervised Learning")
add_bullet("ไม่มี label → โมเดลหา pattern เองจากข้อมูล")
add_bullet("Clustering: จัดกลุ่ม data points ที่คล้ายกันเข้าด้วยกัน")
add_bullet("คุณภาพ clustering ที่ดี: High intra-class similarity + Low inter-class similarity")

add_heading2("9.2 Distance Functions")
add_table(
    ["Distance", "สูตร", "หมายเหตุ"],
    [
        ["Euclidean",  "d(X,Y) = √Σ(xᵢ - yᵢ)²",             "ระยะทางตรง (straight line)"],
        ["Manhattan",  "d(X,Y) = Σ|xᵢ - yᵢ|",               "ระยะทางแบบ city block"],
        ["Cosine",     "d(X,Y) = (X·Y) / (||X|| × ||Y||)",   "วัดมุมระหว่างเวกเตอร์ (text)"],
        ["Hamming",    "จำนวน position ที่แตกต่างกัน",        "ใช้กับข้อมูล categorical/binary"],
    ],
    [3, 6, 5]
)

add_heading2("9.3 Clustering Approaches")
add_heading3("1) Partitioning-based: K-Means")
add_body("แบ่งข้อมูลเป็น k cluster โดยใช้ centroid (จุดศูนย์กลาง)")
add_heading3("K-Means Algorithm Steps:")
add_bullet("ขั้นที่ 1: กำหนด k และสุ่ม initial centroids", level=1)
add_bullet("ขั้นที่ 2: คำนวณ centroid ของแต่ละ cluster", level=1)
add_bullet("ขั้นที่ 3: กำหนด sample แต่ละตัวให้ cluster ที่ centroid ใกล้ที่สุด", level=1)
add_bullet("ขั้นที่ 4: วนกลับขั้นที่ 2 จนกว่าไม่มีการเปลี่ยน assignment", level=1)
add_body("Objective: Ĉ = argmin{SSE(C)}")
add_formula("SSE = Σ_clusters Σ_points dist(point, centroid)²")

add_heading3("การเลือก k")
add_bullet("Elbow Method: Plot WCSS vs k → หาจุด 'elbow'")
add_formula("WCSS = Σ squared distances from all points to their cluster centroid")
add_bullet("Silhouette Score: วัดว่า sample อยู่ใน cluster ที่เหมาะสมแค่ไหน")
add_bullet("  +1 → จัดกลุ่มถูกต้องมาก,  0 → อยู่บน boundary,  -1 → น่าจะอยู่ cluster ผิด")

add_heading3("2) Hierarchical Clustering")
add_bullet("สร้าง nested clusters เป็น tree structure (dendrogram)")
add_bullet("Agglomerative (bottom-up): เริ่มจากแต่ละ sample เป็น cluster → merge ทีละคู่")
add_bullet("Divisive (top-down): เริ่มจาก cluster เดียว → แบ่งออก")
add_body("Agglomerative Algorithm:")
add_bullet("1. คำนวณ proximity matrix", level=1)
add_bullet("2. ให้ทุก data point เป็น cluster ของตัวเอง", level=1)
add_bullet("3. Merge คู่ที่ใกล้ที่สุด", level=1)
add_bullet("4. Update proximity matrix", level=1)
add_bullet("5. ทำซ้ำจนเหลือ 1 cluster", level=1)
add_body("Complete Linkage: MAX((P3,P6),P1) = MAX(P3,P1), (P6,P1)) — ใช้ distance ไกลสุด")

add_heading3("3) Density-based: DBSCAN")
add_bullet("จัดกลุ่มตามความหนาแน่นของข้อมูล (density-based)")
add_bullet("รับมือกับ cluster รูปทรงแปลก (non-convex shapes) ได้ดี")
add_bullet("จัดการ noise/outliers ได้โดยอัตโนมัติ (ไม่ต้องกำหนดจำนวน cluster)")
add_bullet("Parameters: ε (epsilon = radius), MinPts (จำนวน sample ขั้นต่ำ)")
add_table(
    ["Algorithm", "กำหนด k?", "จัดการ outlier?", "รูปทรง cluster"],
    [
        ["K-Means",       "ใช่",  "ไม่",  "Convex (กลม)"],
        ["Hierarchical",  "ไม่",  "บางส่วน", "ยืดหยุ่น"],
        ["DBSCAN",        "ไม่",  "ใช่",  "ทุกรูปทรง"],
    ],
    [4, 3, 3, 4]
)

page_break()

# ════════════════════════════════════════════════════════════════════════════════
# FORMULA CHEAT SHEET
# ════════════════════════════════════════════════════════════════════════════════
add_heading1("ตาราง Cheat Sheet สูตรทั้งหมด (Formula Cheat Sheet)")

add_heading2("บทที่ 6 — Classification Formulas")
add_table(
    ["สูตร/Concept", "นิยาม"],
    [
        ["Entropy(t)",          "-Σⱼ p(j|t) · log₂ p(j|t)"],
        ["Information Gain",    "Entropy(parent) - Σᵢ (nᵢ/n)·Entropy(childᵢ)"],
        ["Gini Index",          "1 - Σⱼ [p(j|t)]²"],
        ["Accuracy",            "(TP+TN) / (TP+TN+FP+FN)"],
        ["Recall",              "TP / (TP+FN)"],
        ["Precision",           "TP / (TP+FP)"],
        ["F1-Score",            "2×(P×R) / (P+R)"],
        ["Logistic (Sigmoid)",  "1 / (1 + e^(-z))"],
    ],
    [6, 8]
)

add_heading2("บทที่ 7 — Regression Formulas")
add_table(
    ["สูตร/Concept", "นิยาม"],
    [
        ["Linear Regression",   "ŷ = w₀ + w₁x₁ + ... + wₙxₙ"],
        ["MSE",                 "(1/n) Σ(yᵢ - ŷᵢ)²"],
        ["RMSE",                "√[(1/n) Σ(yᵢ - ŷᵢ)²]"],
        ["MAE",                 "(1/n) Σ|yᵢ - ŷᵢ|"],
        ["Cost Function J",     "(1/n) Σ(yᵢ - ŷᵢ)²"],
        ["Gradient Descent w",  "w ← w - α · (∂J/∂w)"],
        ["Gradient Descent b",  "b ← b - α · (∂J/∂b)"],
        ["R²",                  "1 - SSE/SSyy  where SSyy=Σ(y-ȳ)², SSE=Σ(y-ŷ)²"],
        ["Ridge Loss (L2)",     "Σ(y-ŷ)² + λΣβⱼ²"],
        ["Lasso Loss (L1)",     "Σ(y-ŷ)² + λΣ|βⱼ|"],
    ],
    [6, 8]
)

add_heading2("บทที่ 8 — Parameter Tuning Concepts")
add_table(
    ["Concept", "สรุป"],
    [
        ["Overfitting",         "High train acc, Low test acc → โมเดลซับซ้อนเกิน"],
        ["Underfitting",        "Low train & test acc → โมเดลเรียบง่ายเกิน"],
        ["k-Fold CV",           "แบ่ง k ส่วน วนทดสอบ k รอบ → ค่าเฉลี่ย accuracy"],
        ["Grid Search",         "ค้นหา exhaustively ทุก combination"],
        ["Random Search",       "สุ่ม sample จาก hyperparameter space"],
        ["Batch GD update/epoch","1 update"],
        ["SGD update/epoch",    "n updates (n = samples)"],
        ["Mini-batch update",   "n/batch_size updates"],
        ["Acc (CM)",            "(TP+TN)/(TP+TN+FP+FN)"],
        ["F1 (CM)",             "2PR/(P+R)"],
    ],
    [5, 9]
)

add_heading2("บทที่ 9 — Clustering Formulas")
add_table(
    ["สูตร/Concept", "นิยาม"],
    [
        ["Euclidean distance",   "d(X,Y) = √Σ(xᵢ-yᵢ)²"],
        ["Manhattan distance",   "d(X,Y) = Σ|xᵢ-yᵢ|"],
        ["Cosine similarity",    "sim(X,Y) = (X·Y)/(||X||·||Y||)"],
        ["WCSS (Inertia)",       "Σ_clusters Σ_points dist(p, centroid)²"],
        ["K-Means objective",    "Ĉ = argmin_C {SSE(C)}"],
        ["Silhouette Score",     "s(i) = (b(i)-a(i)) / max(a(i),b(i))   range: [-1,+1]"],
    ],
    [6, 8]
)

page_break()

# ════════════════════════════════════════════════════════════════════════════════
# EXAM QUESTIONS
# ════════════════════════════════════════════════════════════════════════════════
def add_exam_section(chapter_num, chapter_title):
    add_heading1(f"แนวข้อสอบ — บทที่ {chapter_num}: {chapter_title}")

def add_mcq_section():
    add_heading2("ส่วนที่ 1: ข้อสอบปรนัย (Multiple Choice Questions)")

def add_calc_section():
    add_heading2("ส่วนที่ 2: คำนวณ (Calculation Questions)")

def add_analysis_section():
    add_heading2("ส่วนที่ 3: วิเคราะห์ (Analysis Questions)")

def mcq(num, question, choices, answer, explanation=""):
    p = doc.add_paragraph()
    run = p.add_run(f"ข้อ {num}. {question}")
    set_font(run, 13, bold=True)
    for c in choices:
        add_bullet(c, level=1, size=13)
    p2 = doc.add_paragraph()
    run2 = p2.add_run(f"   ✓ เฉลย: {answer}")
    set_font(run2, 13, color=TEAL)
    if explanation:
        run3 = p2.add_run(f"  ({explanation})")
        set_font(run3, 12, italic=True)

def calc_q(num, question, solution_steps):
    p = doc.add_paragraph()
    run = p.add_run(f"ข้อ {num}. {question}")
    set_font(run, 13, bold=True)
    p2 = doc.add_paragraph()
    run2 = p2.add_run("วิธีทำ:")
    set_font(run2, 13, bold=True, color=TEAL)
    for step in solution_steps:
        add_bullet(step, level=1, size=12)

def analysis_q(num, question, answer_points):
    p = doc.add_paragraph()
    run = p.add_run(f"ข้อ {num}. {question}")
    set_font(run, 13, bold=True)
    p2 = doc.add_paragraph()
    run2 = p2.add_run("แนวคำตอบ:")
    set_font(run2, 13, bold=True, color=TEAL)
    for pt in answer_points:
        add_bullet(pt, level=1, size=12)

# ─── Chapter 6 Exam ────────────────────────────────────────────────────────────
add_exam_section(6, "ML Classification")
add_mcq_section()

mcq(1, "Entropy ของ node ที่มีข้อมูล 2 class เท่ากัน (50:50) มีค่าเท่าใด?",
    ["a) 0", "b) 0.5", "c) 1.0", "d) 2.0"],
    "c) 1.0",
    "Entropy = -(0.5·log₂0.5 + 0.5·log₂0.5) = -(-1) = 1")

mcq(2, "จาก Confusion Matrix: TP=80, TN=70, FP=10, FN=20 Accuracy มีค่าเท่าใด?",
    ["a) 80%", "b) 83.3%", "c) 90%", "d) 75%"],
    "b) 83.3%",
    "(80+70)/(80+70+10+20) = 150/180 = 0.833")

mcq(3, "อัลกอริทึมใดที่ใช้ 'จำนวนเพื่อนบ้านที่ใกล้ที่สุด' เป็น hyperparameter หลัก?",
    ["a) Naive Bayes", "b) Decision Tree", "c) K-Nearest Neighbors", "d) SVM"],
    "c) K-Nearest Neighbors")

mcq(4, "Random Forest แตกต่างจาก Decision Tree อย่างไร?",
    ["a) ใช้ Entropy แทน Gini",
     "b) รวม Decision Trees หลายต้นด้วย Bagging",
     "c) ไม่มี leaf node",
     "d) ใช้ Linear Regression ภายใน"],
    "b) รวม Decision Trees หลายต้นด้วย Bagging")

mcq(5, "Recall = TP/(TP+FN) วัดอะไร?",
    ["a) ความแม่นยำของ prediction ที่เป็น positive",
     "b) สัดส่วนของ positive จริงที่โมเดลตรวจพบได้",
     "c) ความถูกต้องโดยรวม",
     "d) ความสามารถในการแยก class"],
    "b) สัดส่วนของ positive จริงที่โมเดลตรวจพบได้")

mcq(6, "Gini Index = 0 หมายความว่าอะไร?",
    ["a) Node มีความไม่แน่นอนสูงสุด",
     "b) Node เป็น pure (มีแค่ 1 class)",
     "c) ต้องแบ่ง node ต่อ",
     "d) ข้อมูลมี noise มาก"],
    "b) Node เป็น pure (มีแค่ 1 class)")

mcq(7, "SVM มีเป้าหมายหลักคืออะไร?",
    ["a) Minimize entropy",
     "b) หา hyperplane ที่แบ่ง class โดยมี margin สูงสุด",
     "c) สร้าง decision tree ลึกที่สุด",
     "d) คำนวณ probability ของแต่ละ class"],
    "b) หา hyperplane ที่แบ่ง class โดยมี margin สูงสุด")

mcq(8, "Information Gain ของ feature ที่ดีที่สุดใน Decision Tree ควรมีค่า...",
    ["a) ต่ำสุด", "b) เท่ากับ 0", "c) สูงสุด", "d) เท่ากับ 1"],
    "c) สูงสุด",
    "Gain สูงหมายถึง feature นั้นช่วยแบ่งข้อมูลได้ดีที่สุด")

mcq(9, "F1-Score เหมาะใช้เมื่อใด?",
    ["a) ข้อมูล balanced",
     "b) ข้อมูล imbalanced และต้องการ balance precision กับ recall",
     "c) ต้องการ maximize TN",
     "d) ข้อมูลเป็น regression"],
    "b) ข้อมูล imbalanced และต้องการ balance precision กับ recall")

mcq(10, "Logistic Regression ใช้ function ใดแปลงผลลัพธ์เป็นค่าความน่าจะเป็น?",
    ["a) ReLU", "b) Tanh", "c) Sigmoid", "d) Softmax"],
    "c) Sigmoid",
    "P(y=1|x) = 1/(1+e^(-z))")

add_calc_section()

calc_q(1,
    "จงคำนวณ Entropy ของ node ที่มีข้อมูล: Class A = 4 ตัวอย่าง, Class B = 6 ตัวอย่าง",
    ["p(A) = 4/10 = 0.4,  p(B) = 6/10 = 0.6",
     "Entropy = -(0.4·log₂0.4) - (0.6·log₂0.6)",
     "= -(0.4·(-1.322)) - (0.6·(-0.737))",
     "= 0.529 + 0.442",
     "= 0.971"])

calc_q(2,
    "Confusion Matrix: TP=90, TN=85, FP=15, FN=10 จงคำนวณ Accuracy, Precision, Recall, F1",
    ["Accuracy = (90+85)/(90+85+15+10) = 175/200 = 0.875 = 87.5%",
     "Precision = 90/(90+15) = 90/105 = 0.857 = 85.7%",
     "Recall = 90/(90+10) = 90/100 = 0.900 = 90.0%",
     "F1 = 2×(0.857×0.900)/(0.857+0.900) = 2×0.771/1.757 = 0.878 = 87.8%"])

calc_q(3,
    "คำนวณ Information Gain ของ feature 'student' จากข้อมูลต่อไปนี้: Parent Entropy = 0.94, "
    "Child yes (7 samples, Entropy=0.985), Child no (7 samples, Entropy=0.592)",
    ["Gain = 0.94 - [(7/14)×0.985 + (7/14)×0.592]",
     "= 0.94 - [0.5×0.985 + 0.5×0.592]",
     "= 0.94 - [0.4925 + 0.296]",
     "= 0.94 - 0.7885",
     "= 0.152"])

add_analysis_section()

analysis_q(1,
    "อธิบายความแตกต่างระหว่าง Precision กับ Recall และยกตัวอย่างสถานการณ์ที่ควรให้ความสำคัญกับ Recall มากกว่า Precision",
    ["Precision = TP/(TP+FP) — วัดว่า prediction ที่บอกว่า positive นั้น ถูกต้องกี่%",
     "Recall = TP/(TP+FN) — วัดว่า positive จริงๆ ถูกตรวจพบได้กี่%",
     "Recall สำคัญกว่าเมื่อ FN มีต้นทุนสูง เช่น การตรวจโรคมะเร็ง (พลาดผู้ป่วยจริง = อันตราย)",
     "Precision สำคัญกว่าเมื่อ FP มีต้นทุนสูง เช่น spam filter (email สำคัญถูก block = ปัญหา)"])

analysis_q(2,
    "เปรียบเทียบ Decision Tree กับ Random Forest ในด้านการ Overfitting",
    ["Decision Tree: มีแนวโน้ม Overfit สูงหากไม่กำหนด max_depth เพราะต้นไม้เติบโตจนจำ noise",
     "Random Forest: ลด Overfitting ด้วย Bagging — สร้างหลาย trees จาก random subset ของข้อมูลและ features",
     "ผลลัพธ์ใน RF มาจากการ vote ของหลาย trees → ลด variance, เพิ่ม generalization",
     "RF ทนทานต่อ noise กว่า แต่ interpretability ต่ำกว่า Decision Tree เดี่ยว"])

page_break()

# ─── Chapter 7 Exam ────────────────────────────────────────────────────────────
add_exam_section(7, "ML Regression")
add_mcq_section()

mcq(1, "MSE แตกต่างจาก MAE อย่างไร?",
    ["a) MSE ใช้ค่าสัมบูรณ์, MAE ใช้กำลังสอง",
     "b) MSE กำลังสอง error จึงให้น้ำหนัก outlier มากกว่า MAE",
     "c) MSE ใช้เมื่อข้อมูล imbalanced เท่านั้น",
     "d) ไม่มีความแตกต่าง"],
    "b) MSE กำลังสอง error จึงให้น้ำหนัก outlier มากกว่า MAE")

mcq(2, "ใน Gradient Descent: w = w - α·(∂J/∂w) ตัวแปร α หมายถึงอะไร?",
    ["a) Loss function", "b) Regularization term", "c) Learning rate", "d) Number of epochs"],
    "c) Learning rate")

mcq(3, "Ridge Regression แตกต่างจาก Lasso Regression อย่างไร?",
    ["a) Ridge ใช้ |β|, Lasso ใช้ β²",
     "b) Ridge ใช้ β² (L2), Lasso ใช้ |β| (L1) — Lasso อาจทำให้ weight = 0",
     "c) Lasso ใช้ลด Underfitting",
     "d) Ridge เหมาะกับ classification เท่านั้น"],
    "b) Ridge ใช้ β² (L2), Lasso ใช้ |β| (L1) — Lasso อาจทำให้ weight = 0")

mcq(4, "R² = 0.85 หมายความว่าอะไร?",
    ["a) โมเดล error 85%",
     "b) โมเดลอธิบายความแปรปรวนของ y ได้ 85%",
     "c) Accuracy = 85%",
     "d) MSE = 0.15"],
    "b) โมเดลอธิบายความแปรปรวนของ y ได้ 85%")

mcq(5, "Polynomial Regression degree สูงมากๆ ก่อให้เกิดปัญหาใด?",
    ["a) Underfitting", "b) Overfitting", "c) ลด MSE เสมอ", "d) เพิ่ม bias"],
    "b) Overfitting")

mcq(6, "RMSE มีข้อดีอะไรเหนือ MSE?",
    ["a) คำนวณเร็วกว่า",
     "b) มีหน่วยเดียวกับตัวแปร y ทำให้ตีความง่ายกว่า",
     "c) ไม่ sensitive กับ outlier",
     "d) ค่าน้อยกว่า MSE เสมอ"],
    "b) มีหน่วยเดียวกับตัวแปร y ทำให้ตีความง่ายกว่า")

mcq(7, "Decision Tree Regression ต่างจาก Linear Regression อย่างไร?",
    ["a) ไม่ต้องสมมติความสัมพันธ์เชิงเส้น",
     "b) ใช้ Gradient Descent",
     "c) ต้องการข้อมูลมากกว่า",
     "d) ทำงานได้เฉพาะ classification"],
    "a) ไม่ต้องสมมติความสัมพันธ์เชิงเส้น")

mcq(8, "หาก learning rate สูงเกินไป Gradient Descent จะเกิดอะไร?",
    ["a) เรียนช้ามาก", "b) Converge เร็ว", "c) Diverge (ไม่ลู่เข้า)", "d) Underfitting"],
    "c) Diverge (ไม่ลู่เข้า)")

mcq(9, "R² สูงสุดที่เป็นไปได้คือเท่าใด?",
    ["a) 100", "b) 1.0", "c) ∞", "d) 0.999"],
    "b) 1.0",
    "R²=1 → โมเดล predict ได้สมบูรณ์, R²<0 เป็นไปได้หากโมเดลแย่กว่า baseline")

mcq(10, "Lasso มีข้อดีพิเศษอะไรเหนือ Ridge?",
    ["a) ลด MSE มากกว่า",
     "b) ทำ feature selection โดยทำให้ weight บาง feature = 0",
     "c) Training เร็วกว่า",
     "d) ไม่มี hyperparameter"],
    "b) ทำ feature selection โดยทำให้ weight บาง feature = 0")

add_calc_section()

calc_q(1,
    "กำหนดข้อมูล: y = [3, 5, 4, 7, 6], ŷ = [2.5, 5.5, 4.0, 6.5, 5.5] จงคำนวณ MSE และ MAE",
    ["Errors: [0.5, -0.5, 0, 0.5, 0.5]",
     "MSE = (0.5² + 0.5² + 0² + 0.5² + 0.5²) / 5 = (0.25+0.25+0+0.25+0.25)/5 = 1.0/5 = 0.20",
     "MAE = (|0.5|+|0.5|+|0|+|0.5|+|0.5|) / 5 = 2.0/5 = 0.40",
     "RMSE = √0.20 = 0.447"])

calc_q(2,
    "กำหนด SSyy = 200, SSE = 30 จงคำนวณ R²",
    ["R² = 1 - SSE/SSyy",
     "= 1 - 30/200",
     "= 1 - 0.15",
     "= 0.85",
     "→ โมเดลอธิบายความแปรปรวนได้ 85%"])

calc_q(3,
    "โมเดลมี w=2.5, b=1.0, α=0.1, ∂J/∂w=0.4, ∂J/∂b=0.2 จงคำนวณ w และ b ใหม่หลัง 1 step",
    ["w_new = w - α·(∂J/∂w) = 2.5 - 0.1×0.4 = 2.5 - 0.04 = 2.46",
     "b_new = b - α·(∂J/∂b) = 1.0 - 0.1×0.2 = 1.0 - 0.02 = 0.98"])

add_analysis_section()

analysis_q(1,
    "อธิบายปัญหา Overfitting ใน Regression และวิธีใช้ Regularization แก้ปัญหา",
    ["Overfitting ใน Regression: โมเดลที่ degree สูงมากจะผ่านทุก data point แต่ predict ข้อมูลใหม่ได้ไม่ดี",
     "Ridge (L2): เพิ่ม λΣβⱼ² ใน loss → บังคับให้ weight มีค่าน้อยลง แต่ไม่เป็น 0",
     "Lasso (L1): เพิ่ม λΣ|βⱼ| ใน loss → บาง feature มี weight = 0 (feature selection)",
     "λ ใหญ่ = regularization แรง (ลด Overfitting แต่เสี่ยง Underfitting)",
     "ควรใช้ Cross-Validation เพื่อหาค่า λ ที่เหมาะสม"])

analysis_q(2,
    "เปรียบเทียบ Batch Gradient Descent, Stochastic GD และ Mini-batch GD",
    ["Batch GD: ใช้ทั้ง dataset คำนวณ gradient → เสถียร แต่ช้าและใช้ memory มาก",
     "Stochastic GD: ใช้ 1 sample → เร็ว แต่ noisy อาจไม่ converge",
     "Mini-batch GD: ใช้ batch 32-64 samples → สมดุลระหว่าง 2 แบบ (ใช้ใน Deep Learning)",
     "ในทางปฏิบัติ Mini-batch GD นิยมใช้มากที่สุด"])

page_break()

# ─── Chapter 8 Exam ────────────────────────────────────────────────────────────
add_exam_section(8, "Model Parameter Tuning")
add_mcq_section()

mcq(1, "Hyperparameter ต่างจาก Learnable Parameter อย่างไร?",
    ["a) Hyperparameter ถูกเรียนรู้จาก data โดยตรง",
     "b) Hyperparameter กำหนดโดยผู้ใช้ก่อน training (ไม่ได้ optimize ระหว่าง training)",
     "c) Hyperparameter ใช้เฉพาะกับ Neural Networks",
     "d) Learnable Parameter ต้อง tune ด้วย Grid Search"],
    "b) Hyperparameter กำหนดโดยผู้ใช้ก่อน training (ไม่ได้ optimize ระหว่าง training)")

mcq(2, "max_depth ใน Decision Tree มีผลอย่างไร?",
    ["a) กำหนดจำนวน feature ที่ใช้",
     "b) ควบคุมความลึกสูงสุด — ค่าน้อย = ป้องกัน Overfitting",
     "c) กำหนด learning rate",
     "d) เลือก Gini หรือ Entropy"],
    "b) ควบคุมความลึกสูงสุด — ค่าน้อย = ป้องกัน Overfitting")

mcq(3, "Stratified k-Fold Cross-Validation เหมาะกับสถานการณ์ใด?",
    ["a) ข้อมูล time-series",
     "b) ข้อมูล balanced",
     "c) ข้อมูล imbalanced (class distribution ไม่เท่ากัน)",
     "d) ข้อมูลขนาดใหญ่มาก"],
    "c) ข้อมูล imbalanced (class distribution ไม่เท่ากัน)")

mcq(4, "Grid Search ต่างจาก Random Search อย่างไร?",
    ["a) Grid Search สุ่ม, Random Search ค้นหาทุก combination",
     "b) Grid Search ค้นหาทุก combination, Random Search สุ่ม sample",
     "c) Grid Search เร็วกว่า",
     "d) Random Search รับประกัน global optimal"],
    "b) Grid Search ค้นหาทุก combination, Random Search สุ่ม sample")

mcq(5, "Validation Set ต่างจาก Test Set อย่างไร?",
    ["a) ไม่ต่างกัน",
     "b) Validation ใช้ระหว่าง training เพื่อ tune HP; Test ใช้ครั้งเดียวหลัง training",
     "c) Test ใช้สำหรับ tune hyperparameters",
     "d) Validation เป็นส่วนหนึ่งของ Test Set"],
    "b) Validation ใช้ระหว่าง training เพื่อ tune HP; Test ใช้ครั้งเดียวหลัง training")

mcq(6, "Leave-One-Out Cross-Validation (LOO-CV) มีข้อเสียใด?",
    ["a) ไม่สามารถใช้กับ imbalanced data",
     "b) Computationally expensive เพราะต้อง train n รอบ",
     "c) ใช้ได้เฉพาะ regression",
     "d) ให้ผล biased estimate"],
    "b) Computationally expensive เพราะต้อง train n รอบ")

mcq(7, "Dropout ใช้แก้ปัญหาใด?",
    ["a) Underfitting ใน Decision Tree",
     "b) Overfitting ใน Neural Networks โดย random drop connections",
     "c) Local Minima ใน Gradient Descent",
     "d) Class imbalance ใน Classification"],
    "b) Overfitting ใน Neural Networks โดย random drop connections")

mcq(8, "min_samples_split = 10 หมายความว่า...",
    ["a) ต้นไม้มีได้ไม่เกิน 10 leaf nodes",
     "b) node ที่มีน้อยกว่า 10 samples จะไม่ถูกแบ่งต่อ",
     "c) แต่ละ leaf ต้องมีอย่างน้อย 10 samples",
     "d) จำนวน features สูงสุดที่พิจารณาคือ 10"],
    "b) node ที่มีน้อยกว่า 10 samples จะไม่ถูกแบ่งต่อ")

mcq(9, "Local Minima ในการ optimize hyperparameter คืออะไร?",
    ["a) จุดที่ loss สูงสุด",
     "b) จุดที่ loss ต่ำกว่าบริเวณใกล้เคียงแต่ไม่ใช่ต่ำสุดทั่วโลก (global minimum)",
     "c) จุดที่ learning rate = 0",
     "d) จุดที่ model เริ่มต้น training"],
    "b) จุดที่ loss ต่ำกว่าบริเวณใกล้เคียงแต่ไม่ใช่ต่ำสุดทั่วโลก (global minimum)")

mcq(10, "k-Fold Cross-Validation ที่ k=5 หมายความว่า...",
    ["a) train 5 models แยกกัน",
     "b) แบ่งข้อมูล 5 ส่วน วนทดสอบ 5 รอบ แต่ละรอบใช้ 1 ส่วนเป็น test",
     "c) ใช้ข้อมูล 5% เป็น test",
     "d) train 5 epochs"],
    "b) แบ่งข้อมูล 5 ส่วน วนทดสอบ 5 รอบ แต่ละรอบใช้ 1 ส่วนเป็น test")

add_calc_section()

calc_q(1,
    "Grid Search สำหรับ Decision Tree: max_depth={2,3,4}, criterion={gini,entropy}, "
    "min_samples_split={2,5,10} มีจำนวน combination ทั้งหมดเท่าใด?",
    ["จำนวน max_depth = 3 ค่า",
     "จำนวน criterion = 2 ค่า",
     "จำนวน min_samples_split = 3 ค่า",
     "Total combinations = 3 × 2 × 3 = 18 combinations"])

calc_q(2,
    "5-Fold CV ให้ accuracy: [0.82, 0.85, 0.80, 0.88, 0.83] จงหา mean accuracy และ std deviation",
    ["Mean = (0.82+0.85+0.80+0.88+0.83)/5 = 4.18/5 = 0.836 = 83.6%",
     "Variance = Σ(xᵢ-μ)²/n = [(0.82-0.836)²+(0.85-0.836)²+(0.80-0.836)²+(0.88-0.836)²+(0.83-0.836)²]/5",
     "= [0.000256+0.000196+0.001296+0.001936+0.000036]/5 = 0.00372/5 = 0.000744",
     "Std = √0.000744 ≈ 0.027 = 2.7%"])

add_analysis_section()

analysis_q(1,
    "อธิบาย Bias-Variance Tradeoff และความสัมพันธ์กับ Overfitting/Underfitting",
    ["Bias = ความผิดพลาดจากการสมมติที่ผิด (โมเดลง่ายเกินไป) → Underfitting",
     "Variance = ความผิดพลาดจากความ sensitive ต่อ noise (โมเดลซับซ้อนเกิน) → Overfitting",
     "Trade-off: ลด Bias มักเพิ่ม Variance และกลับกัน",
     "Cross-Validation ช่วยหา sweet spot ระหว่าง bias และ variance",
     "Regularization ช่วย balance โดยเพิ่ม constraint บน model complexity"])

page_break()

# ─── Chapter 9 Exam ────────────────────────────────────────────────────────────
add_exam_section(9, "ML Clustering")
add_mcq_section()

mcq(1, "K-Means Clustering กำหนดค่า k หมายถึงอะไร?",
    ["a) จำนวน iterations", "b) จำนวน clusters ที่ต้องการ",
     "c) จำนวน features", "d) จำนวน data points"],
    "b) จำนวน clusters ที่ต้องการ")

mcq(2, "Elbow Method ใช้วิเคราะห์อะไร?",
    ["a) Silhouette Score vs k",
     "b) WCSS vs จำนวน clusters เพื่อหา optimal k",
     "c) Distance function ที่เหมาะสม",
     "d) จำนวน iterations ที่ต้องการ"],
    "b) WCSS vs จำนวน clusters เพื่อหา optimal k")

mcq(3, "Silhouette Score = +1 หมายความว่าอะไร?",
    ["a) sample อยู่บน boundary ระหว่าง 2 clusters",
     "b) sample อยู่ใน cluster ที่เหมาะสมมาก (ห่างจาก cluster อื่นๆ)",
     "c) sample น่าจะอยู่ผิด cluster",
     "d) ข้อมูลเป็น outlier"],
    "b) sample อยู่ใน cluster ที่เหมาะสมมาก (ห่างจาก cluster อื่นๆ)")

mcq(4, "Euclidean distance ระหว่าง A(1,2) และ B(4,6) มีค่าเท่าใด?",
    ["a) 3", "b) 4", "c) 5", "d) 7"],
    "c) 5",
    "d = √((4-1)²+(6-2)²) = √(9+16) = √25 = 5")

mcq(5, "DBSCAN เหนือกว่า K-Means ในด้านใด?",
    ["a) เร็วกว่า",
     "b) จัดการ outlier และ cluster รูปทรงแปลกได้ดีกว่า",
     "c) ไม่ต้องกำหนด parameter ใดเลย",
     "d) ใช้ Euclidean distance เสมอ"],
    "b) จัดการ outlier และ cluster รูปทรงแปลกได้ดีกว่า")

mcq(6, "Agglomerative Hierarchical Clustering เริ่มต้นอย่างไร?",
    ["a) เริ่มด้วย cluster เดียวแล้วแบ่งออก",
     "b) เริ่มด้วยแต่ละ data point เป็น cluster แล้ว merge",
     "c) สุ่ม centroids แล้ว assign",
     "d) สร้าง dendrogram จากบนลงล่าง"],
    "b) เริ่มด้วยแต่ละ data point เป็น cluster แล้ว merge")

mcq(7, "Manhattan distance ระหว่าง A(1,3) และ B(4,7) คือเท่าใด?",
    ["a) 5", "b) 7", "c) 3", "d) 4"],
    "b) 7",
    "d = |4-1| + |7-3| = 3 + 4 = 7")

mcq(8, "Clustering เป็น Unsupervised Learning เพราะ...",
    ["a) ใช้ labeled data",
     "b) ไม่มี ground truth labels — โมเดลหา pattern เองจากข้อมูล",
     "c) ใช้ loss function เหมือน Regression",
     "d) ต้องการ validation set"],
    "b) ไม่มี ground truth labels — โมเดลหา pattern เองจากข้อมูล")

mcq(9, "Dendrogram ใช้กับ clustering approach ใด?",
    ["a) K-Means", "b) DBSCAN", "c) Hierarchical Clustering", "d) K-Medoids"],
    "c) Hierarchical Clustering")

mcq(10, "Cosine Similarity เหมาะกับการวัด distance ของข้อมูลประเภทใดที่สุด?",
    ["a) ข้อมูลตัวเลขต่อเนื่อง",
     "b) ข้อมูล text/document vectors",
     "c) ข้อมูล binary categorical",
     "d) ข้อมูล time-series"],
    "b) ข้อมูล text/document vectors")

add_calc_section()

calc_q(1,
    "กำหนดข้อมูล 2 มิติ: A(1,2), B(1,4), C(1,0), D(10,2), E(10,4), F(10,0) "
    "สมมติ k=2 Cluster1={A,B,C}, Cluster2={D,E,F} จงคำนวณ WCSS",
    ["Centroid1 = ((1+1+1)/3, (2+4+0)/3) = (1, 2)",
     "Centroid2 = ((10+10+10)/3, (2+4+0)/3) = (10, 2)",
     "WCSS₁: A→0, B→(0²+2²)=4, C→(0²+2²)=4  →  WCSS₁ = 8",
     "WCSS₂: D→0, E→(0²+2²)=4, F→(0²+2²)=4  →  WCSS₂ = 8",
     "WCSS_Total = 8 + 8 = 16"])

calc_q(2,
    "คำนวณ Euclidean distance ระหว่าง P1(0.4,0.53) กับ P3(0.35,0.32)",
    ["d(P1,P3) = √((0.4-0.35)² + (0.53-0.32)²)",
     "= √((0.05)² + (0.21)²)",
     "= √(0.0025 + 0.0441)",
     "= √0.0466",
     "≈ 0.216"])

calc_q(3,
    "Agglomerative Clustering: จาก proximity matrix ค้นหาคู่ที่ใกล้ที่สุดใน "
    "{P1,P2,P3,P4,P5,P6} กำหนด distance matrix บางส่วน: d(P3,P6)=0.10, d(P2,P3)=0.14",
    ["เปรียบเทียบ distance ทุกคู่จาก proximity matrix",
     "คู่ที่ใกล้สุดคือ (P3, P6) = 0.10",
     "Merge P3 และ P6 เป็น cluster {P3,P6} ก่อน",
     "อัปเดต proximity matrix: d({P3,P6},P1) = MAX(d(P3,P1), d(P6,P1)) = MAX(0.22, 0.24) = 0.24"])

add_analysis_section()

analysis_q(1,
    "เปรียบเทียบ K-Means, Hierarchical Clustering และ DBSCAN ในด้านข้อดีข้อเสีย",
    ["K-Means: ข้อดี=เร็ว, เหมาะกับ convex clusters; ข้อเสีย=ต้องกำหนด k, sensitive ต่อ outlier",
     "Hierarchical: ข้อดี=ไม่ต้องกำหนด k ล่วงหน้า, สร้าง dendrogram; ข้อเสีย=ช้า O(n²)",
     "DBSCAN: ข้อดี=จัดการ outlier, cluster รูปทรงใดก็ได้; ข้อเสีย=sensitive ต่อ ε และ MinPts",
     "เลือก K-Means เมื่อรู้จำนวน cluster และ cluster กลมๆ",
     "เลือก DBSCAN เมื่อมี outlier หรือ cluster รูปทรงแปลก"])

analysis_q(2,
    "อธิบายปัญหาที่อาจเกิดขึ้นใน K-Means Clustering และวิธีแก้",
    ["ปัญหา 1: Sensitive ต่อ initial centroids → แก้โดยใช้ K-Means++ initialization",
     "ปัญหา 2: ต้องกำหนด k → แก้โดยใช้ Elbow Method หรือ Silhouette Score",
     "ปัญหา 3: ไม่รับมือ outlier ได้ดี → แก้โดยใช้ K-Medoids แทน",
     "ปัญหา 4: เหมาะกับ convex cluster เท่านั้น → แก้โดยใช้ DBSCAN สำหรับ non-convex",
     "ปัญหา 5: WCSS ลดเสมอเมื่อ k เพิ่ม → ใช้ Elbow Point หาจุดที่ลดอย่างมีนัย"])

# ─── Final Save ────────────────────────────────────────────────────────────────
output_path = r"c:\Users\CPE KMUTT\Documents\GitHub\cpe232-datamodel-2025\CPE232_FinalExam_Summary.docx"
doc.save(output_path)
print(f"[OK] Document saved: {output_path}")
